"""
parser.py — DOCX Question Bank Parser.

Reads a .docx file containing a single-table Question Bank and extracts:
  - Metadata (SU00, SU01, BR00, YR00, SE00)
  - Course Outcomes (CO1–CO6)
  - Part A questions grouped by question number
  - Part B questions grouped by question number and sub-part (a/b)
  - Part C questions (same structure as Part B)
  - Marks distribution from section headers
  - Inline mathematical equations (modern OMML and legacy Equation Editor 3.0 / MathType)
  - Embedded diagrams and images
"""

import os
import re
import struct
import xml.etree.ElementTree as ET
from io import BytesIO
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from PIL import Image as PILImage, ImageDraw, ImageFont, ImageChops
from docx import Document


# Mappings for Branch names to Short Branch Code
BRANCH_MAP = {
    'ARTIFICIAL INTELLIGENCE AND DATA SCIENCE': 'ADS',
    'AI & DS': 'ADS',
    'AI AND DS': 'ADS',
    'ADS': 'ADS',
    'COMPUTER SCIENCE AND ENGINEERING': 'CSE',
    'CSE': 'CSE',
    'INFORMATION TECHNOLOGY': 'IT',
    'IT': 'IT',
    'COMPUTER SCIENCE AND BUSINESS SYSTEMS': 'CSBS',
    'CSBS': 'CSBS',
    'ELECTRONICS AND COMMUNICATION ENGINEERING': 'ECE',
    'ECE': 'ECE',
    'ELECTRICAL AND ELECTRONICS ENGINEERING': 'EEE',
    'EEE': 'EEE',
    'MECHANICAL ENGINEERING': 'MECH',
    'MECH': 'MECH',
    'CIVIL ENGINEERING': 'CIVIL',
    'CIVIL': 'CIVIL',
}

# Semester to Roman Numeral & Year calculation
SEM_MAP = {
    '1': ('I', 'I'),
    '1ST': ('I', 'I'),
    'FIRST': ('I', 'I'),
    'I': ('I', 'I'),
    '2': ('II', 'I'),
    '2ND': ('II', 'I'),
    'SECOND': ('II', 'I'),
    'II': ('II', 'I'),
    '3': ('III', 'II'),
    '3RD': ('III', 'II'),
    'THIRD': ('III', 'II'),
    'III': ('III', 'II'),
    '4': ('IV', 'II'),
    '4TH': ('IV', 'II'),
    'FOURTH': ('IV', 'II'),
    'IV': ('IV', 'II'),
    '5': ('V', 'III'),
    '5TH': ('V', 'III'),
    'FIFTH': ('V', 'III'),
    'V': ('V', 'III'),
    '6': ('VI', 'III'),
    '6TH': ('VI', 'III'),
    'SIXTH': ('VI', 'III'),
    'VI': ('VI', 'III'),
    '7': ('VII', 'IV'),
    '7TH': ('VII', 'IV'),
    'SEVENTH': ('VII', 'IV'),
    'VII': ('VII', 'IV'),
    '8': ('VIII', 'IV'),
    '8TH': ('VIII', 'IV'),
    'EIGHTH': ('VIII', 'IV'),
    'VIII': ('VIII', 'IV'),
}


def parse_question_bank(filepath, file_id=None):
    """
    Parse a Question Bank .docx file.
    
    Args:
        filepath: Path to the .docx file.
        file_id: Optional unique identifier for storing extracted images.
    
    Returns:
        dict with metadata, course outcomes, and question sections with associated images and equations.
    """
    if not file_id:
        file_id = os.path.splitext(os.path.basename(filepath))[0]

    doc = Document(filepath)

    if not doc.tables:
        raise ValueError("No tables found in the document. Expected a table-based Question Bank.")

    # Image and equation storage directories for this file
    base_dir = os.path.dirname(os.path.abspath(__file__))
    img_dir = os.path.join(base_dir, 'uploads', 'images', file_id)
    eq_dir = os.path.join(base_dir, 'uploads', 'equations', file_id)
    os.makedirs(img_dir, exist_ok=True)
    os.makedirs(eq_dir, exist_ok=True)

    table = doc.tables[0]
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    # ── Extract Metadata ──────────────────────────────────
    metadata = _extract_metadata(rows)

    # ── Extract Course Outcomes ───────────────────────────
    course_outcomes = _extract_course_outcomes(rows)

    # ── Find section boundaries ───────────────────────────
    part_a_start, part_a_config = None, ''
    part_b_start, part_b_config = None, ''
    part_c_start, part_c_config = None, ''

    for i, row in enumerate(rows):
        text = row[2].lower() if len(row) > 2 else ''
        if not text:
            text = ' '.join(c for c in row if c).lower()

        if re.search(r'part[\s\-\u2013\u2014\ufffd–—]*a\b', text):
            part_a_start = i
            part_a_config = _extract_marks_config(text)
        elif re.search(r'part[\s\-\u2013\u2014\ufffd–—]*b\b', text):
            part_b_start = i
            part_b_config = _extract_marks_config(text)
        elif re.search(r'part[\s\-\u2013\u2014\ufffd–—]*c\b', text):
            part_c_start = i
            part_c_config = _extract_marks_config(text)

    # ── Parse Part A ──────────────────────────────────────
    part_a_end = part_b_start if part_b_start else len(rows)
    part_a_questions = _parse_part_a(table, part_a_start, part_a_end, doc, file_id, img_dir, eq_dir)

    # ── Parse Part B ──────────────────────────────────────
    part_b_end = part_c_start if part_c_start else len(rows)
    part_b_questions = _parse_part_bc(table, part_b_start, part_b_end, doc, file_id, img_dir, eq_dir)

    # ── Parse Part C ──────────────────────────────────────
    part_c_questions = _parse_part_bc(table, part_c_start, len(rows), doc, file_id, img_dir, eq_dir)

    # ── Determine marks per question ──────────────────────
    part_a_marks = _extract_marks_per_question(part_a_config)
    part_b_marks = _extract_marks_per_question(part_b_config)
    part_c_marks = _extract_marks_per_question(part_c_config)

    return {
        'metadata': metadata,
        'course_outcomes': course_outcomes,
        'part_a': {
            'config': part_a_config,
            'marks_per_question': part_a_marks,
            'questions': part_a_questions,
        },
        'part_b': {
            'config': part_b_config,
            'marks_per_question': part_b_marks,
            'questions': part_b_questions,
        },
        'part_c': {
            'config': part_c_config,
            'marks_per_question': part_c_marks,
            'questions': part_c_questions,
        },
    }


def _extract_row_equation_rids(row):
    """
    Find all relationship IDs that belong to legacy equation objects in a row.
    These IDs should NOT be treated as regular question diagrams.
    """
    equation_rids = set()
    for obj in row._element.xpath('.//w:object'):
        ole = obj.find('.//{urn:schemas-microsoft-com:office:office}OLEObject')
        imgd = obj.find('.//{urn:schemas-microsoft-com:vml}imagedata')
        prog = ole.attrib.get('ProgID', '') if ole is not None else ''
        if 'equation' in prog.lower() or 'mathtype' in prog.lower() or 'dsmt' in prog.lower() or imgd is not None:
            if imgd is not None:
                for k, v in imgd.attrib.items():
                    if 'id' in k.lower():
                        equation_rids.add(v)
            if ole is not None:
                for k, v in ole.attrib.items():
                    if 'id' in k.lower():
                        equation_rids.add(v)
    return equation_rids


def _extract_row_images(row, r_idx, doc, file_id, img_dir, equation_rids=None):
    """
    Extract embedded normal question images from a table row XML and save to disk.
    Explicitly excludes equation objects so normal question images remain distinct.
    """
    if equation_rids is None:
        equation_rids = _extract_row_equation_rids(row)

    xml_str = row._element.xml
    rids = re.findall(r'(?:r:embed|r:id|r:link)="([^"]+)"', xml_str)
    extracted_urls = []
    for rid in rids:
        if rid in equation_rids:
            continue
        if rid in doc.part.rels and 'image' in doc.part.rels[rid].target_ref.lower():
            part = doc.part.rels[rid].target_part
            ext = os.path.splitext(part.filename)[1] or '.png'
            filename = f"img_{r_idx}_{rid}{ext}"
            disk_path = os.path.join(img_dir, filename)
            if not os.path.exists(disk_path):
                with open(disk_path, 'wb') as f:
                    f.write(part.blob)
            url = f"/api/images/{file_id}/{filename}"
            if url not in extracted_urls:
                extracted_urls.append(url)
    return extracted_urls


def _extract_metadata(rows):
    """
    Extract metadata fields from QB:
      SU00 = Subject Code (e.g. AIT519)
      SU01 = Subject Name (e.g. Artificial Intelligence)
      BR00 = Short Branch Code (CSE / IT / ADS / CSBS)
      YR00 = Year (I / II / III / IV)
      SE00 = Semester (I / II / III / IV / V / VI / VII / VIII)
    """
    metadata = {
        'su00': '-',
        'su01': '-',
        'br00': '-',
        'yr00': '-',
        'se00': '-',
        'subject_code': '-',
        'subject_name': '-',
        'branch': '-',
        'semester': '-',
        'branch_info': '-',
    }

    raw_br_list = []
    raw_se = '-'
    raw_yr = '-'

    for row in rows[:15]:
        col0 = row[0].strip().upper() if row[0] else ''
        col2 = row[2].strip() if len(row) > 2 else ''

        if not col2:
            continue

        if col0 == 'SE00' or col0.startswith('SE'):
            raw_se = col2

        elif col0.startswith('BR'):
            if col2 not in raw_br_list:
                raw_br_list.append(col2)

        elif col0 == 'YR00' or col0.startswith('YR'):
            raw_yr = col2

        elif col0 == 'SU00' or col0.startswith('SU00'):
            match = re.match(r'([A-Z0-9]+)\s*[–\-]\s*(.*)', col2)
            if match:
                metadata['su00'] = match.group(1).strip()
                if metadata['su01'] == '-':
                    metadata['su01'] = match.group(2).strip()
            else:
                metadata['su00'] = col2

        elif col0 == 'SU01' or col0.startswith('SU01'):
            metadata['su01'] = col2

    # ── Process BR00 (Complete Branch Information) ─────────
    if raw_br_list:
        full_br = ", ".join(raw_br_list)
        metadata['br00'] = full_br
        metadata['branch'] = full_br

    # ── Process SE00 (Semester) & YR00 (Year) ────────────
    if raw_se != '-':
        match_sem = re.search(r'\b(1ST|2ND|3RD|4TH|5TH|6TH|7TH|8TH|1|2|3|4|5|6|7|8|VIII|VII|VI|V|IV|III|II|I)\b', raw_se.upper())
        if match_sem:
            sem_key = match_sem.group(1)
            sem_roman, calc_yr = SEM_MAP.get(sem_key, (sem_key, '-'))
            metadata['se00'] = sem_roman
            metadata['semester'] = sem_roman
            if raw_yr == '-':
                raw_yr = calc_yr

    # ── Process YR00 (Year) ───────────────────────────────
    if raw_yr != '-':
        cleaned_yr = raw_yr.upper()
        match_yr = re.search(r'\b(1ST|2ND|3RD|4TH|1|2|3|4|IV|III|II|I)\b', cleaned_yr)
        if match_yr:
            yr_key = match_yr.group(1)
            yr_roman, _ = SEM_MAP.get(yr_key, (yr_key, '-'))
            metadata['yr00'] = yr_roman
        else:
            metadata['yr00'] = raw_yr

    # Fallbacks for subject code / subject name
    if metadata['su00'] != '-':
        metadata['subject_code'] = metadata['su00']
    if metadata['su01'] != '-':
        metadata['subject_name'] = metadata['su01']

    # Dynamically build Branch / Year / Sem: BR00 / YR00 / SE00
    metadata['branch_info'] = f"{metadata['br00']} / {metadata['yr00']} / {metadata['se00']}"

    return metadata


def _extract_course_outcomes(rows):
    """Extract course outcomes (CO1–CO6) from the metadata section."""
    outcomes = []
    for row in rows[:15]:
        col0 = row[0].strip().upper() if row[0] else ''
        col2 = row[2].strip() if len(row) > 2 else ''

        match = re.match(r'(CO\d+)#?', col0)
        if match and col2:
            outcomes.append({
                'id': match.group(1),
                'text': col2,
            })

    return outcomes


def _extract_marks_config(text):
    """Extract string like '5 x 2 = 10 Marks' from section header text."""
    match = re.search(r'\(\s*\d+\s*[xX*]\s*\d+[^)]*\)', text)
    if match:
        return match.group(0).strip('()').strip()
    return ''


def _extract_marks_per_question(config):
    """Parse mark per question from config string like '5 x 2 = 10 Marks'."""
    match = re.search(r'\d+\s*[xX*]\s*(\d+)', config)
    if match:
        return int(match.group(1))
    return 0


def _parse_shape_style_dimensions(style_str):
    """Extract width_pt and height_pt from shape style string like 'width:120.75pt;height:18pt'."""
    w_pt, h_pt = None, None
    if not style_str:
        return w_pt, h_pt

    w_match = re.search(r'width\s*:\s*([\d.]+)\s*(pt|in|cm|mm|px)?', style_str, re.IGNORECASE)
    if w_match:
        val = float(w_match.group(1))
        unit = (w_match.group(2) or 'pt').lower()
        if unit == 'in': w_pt = val * 72.0
        elif unit == 'cm': w_pt = val * 28.3465
        elif unit == 'mm': w_pt = val * 2.83465
        elif unit == 'px': w_pt = val * 0.75
        else: w_pt = val

    h_match = re.search(r'height\s*:\s*([\d.]+)\s*(pt|in|cm|mm|px)?', style_str, re.IGNORECASE)
    if h_match:
        val = float(h_match.group(1))
        unit = (h_match.group(2) or 'pt').lower()
        if unit == 'in': h_pt = val * 72.0
        elif unit == 'cm': h_pt = val * 28.3465
        elif unit == 'mm': h_pt = val * 2.83465
        elif unit == 'px': h_pt = val * 0.75
        else: h_pt = val

    return w_pt, h_pt


import ctypes
from ctypes import wintypes
from PIL import Image as PILImage, ImageChops

# ── Windows GDI+ High-Fidelity WMF/EMF Vector Rasterizer ────────
_gdiplus = None
_gdiplus_available = False

try:
    if os.name == 'nt':
        _gdiplus = ctypes.windll.gdiplus
        class _GdiplusStartupInput(ctypes.Structure):
            _fields_ = [
                ('GdiplusVersion', wintypes.UINT),
                ('DebugEventCallback', ctypes.c_void_p),
                ('SuppressBackgroundThread', wintypes.BOOL),
                ('SuppressExternalCodecs', wintypes.BOOL)
            ]
        class _CLSID(ctypes.Structure):
            _fields_ = [
                ('Data1', wintypes.DWORD),
                ('Data2', wintypes.WORD),
                ('Data3', wintypes.WORD),
                ('Data4', ctypes.c_ubyte * 8)
            ]
        _token = ctypes.c_ulong()
        _startup_in = _GdiplusStartupInput(1, None, False, False)
        _st = _gdiplus.GdiplusStartup(ctypes.byref(_token), ctypes.byref(_startup_in), None)
        if _st == 0:
            _gdiplus.GdipLoadImageFromFile.argtypes = [ctypes.c_wchar_p, ctypes.POINTER(ctypes.c_void_p)]
            _gdiplus.GdipGetImageDimension.argtypes = [ctypes.c_void_p, ctypes.POINTER(ctypes.c_float), ctypes.POINTER(ctypes.c_float)]
            _gdiplus.GdipCreateBitmapFromScan0.argtypes = [ctypes.c_int, ctypes.c_int, ctypes.c_int, ctypes.c_int, ctypes.c_void_p, ctypes.POINTER(ctypes.c_void_p)]
            _gdiplus.GdipGetImageGraphicsContext.argtypes = [ctypes.c_void_p, ctypes.POINTER(ctypes.c_void_p)]
            _gdiplus.GdipGraphicsClear.argtypes = [ctypes.c_void_p, ctypes.c_uint]
            _gdiplus.GdipSetSmoothingMode.argtypes = [ctypes.c_void_p, ctypes.c_int]
            _gdiplus.GdipSetInterpolationMode.argtypes = [ctypes.c_void_p, ctypes.c_int]
            _gdiplus.GdipSetTextRenderingHint.argtypes = [ctypes.c_void_p, ctypes.c_int]
            _gdiplus.GdipDrawImageRectRect.argtypes = [
                ctypes.c_void_p, ctypes.c_void_p,
                ctypes.c_float, ctypes.c_float, ctypes.c_float, ctypes.c_float,
                ctypes.c_float, ctypes.c_float, ctypes.c_float, ctypes.c_float,
                ctypes.c_int, ctypes.c_void_p, ctypes.c_void_p, ctypes.c_void_p
            ]
            _gdiplus.GdipSaveImageToFile.argtypes = [ctypes.c_void_p, ctypes.c_wchar_p, ctypes.POINTER(_CLSID), ctypes.c_void_p]
            _gdiplus.GdipDeleteGraphics.argtypes = [ctypes.c_void_p]
            _gdiplus.GdipDisposeImage.argtypes = [ctypes.c_void_p]
            _png_clsid = _CLSID(0x557CF406, 0x1A04, 0x11D3, (ctypes.c_ubyte * 8)(0x9A, 0x73, 0x00, 0x00, 0xF8, 0x1E, 0xF3, 0x2E))
            _gdiplus_available = True
except Exception as e:
    _gdiplus_available = False


def _trim_white_borders(img, padding=2):
    """Trim excess whitespace around rendered equation image."""
    try:
        if img.mode != 'RGB':
            img = img.convert('RGB')
        bg = PILImage.new('RGB', img.size, (255, 255, 255))
        diff = ImageChops.difference(img, bg)
        bbox = diff.getbbox()
        if bbox:
            left = max(0, bbox[0] - padding)
            top = max(0, bbox[1] - padding)
            right = min(img.width, bbox[2] + padding)
            bottom = min(img.height, bbox[3] + padding)
            return img.crop((left, top, right, bottom))
    except Exception:
        pass
    return img


def _render_wmf_to_pil(wmf_blob, target_width_pt=None, target_height_pt=None):
    """
    Robust pure-Python WMF (Windows Metafile) rasterizer using Pillow.
    Decodes GDI objects (fonts, pens, brushes) and vector/text primitives (EXTTEXTOUT, TEXTOUT,
    RECTANGLE, POLYGON, POLYPOLYGON, POLYLINE, MOVETO, LINETO).
    Works 100% identically on Linux (Render) and Windows with zero external C dependencies.
    """
    if not wmf_blob or len(wmf_blob) < 40:
        return None

    try:
        offset = 0
        key = struct.unpack('<I', wmf_blob[:4])[0]
        bbox_l, bbox_t, bbox_r, bbox_b, inch = 0, 0, 0, 0, 1440

        if key == 0x9ac6cdd7:  # Placeable WMF header
            offset = 22
            bbox_l, bbox_t, bbox_r, bbox_b, inch = struct.unpack('<hhhhH', wmf_blob[6:16])

        if offset + 18 > len(wmf_blob):
            return None

        file_type, header_size, version, file_size, num_objects, max_record, num_members = struct.unpack(
            '<HHHIHIH', wmf_blob[offset:offset+18]
        )

        rec_offset = offset + 18
        records = []
        while rec_offset < len(wmf_blob):
            if rec_offset + 6 > len(wmf_blob):
                break
            rec_size, rec_func = struct.unpack('<IH', wmf_blob[rec_offset:rec_offset+6])
            if rec_size == 0:
                break
            rec_data = wmf_blob[rec_offset+6 : rec_offset+rec_size*2]
            records.append((rec_func, rec_data))
            rec_offset += rec_size * 2

        # Coordinate extents
        win_org_x, win_org_y = bbox_l, bbox_t
        win_ext_w = (bbox_r - bbox_l) if bbox_r > bbox_l else 1000
        win_ext_h = (bbox_b - bbox_t) if bbox_b > bbox_t else 1000

        for func, data in records:
            if func == 0x020b and len(data) >= 4:  # SETWINDOWORG
                y, x = struct.unpack('<hh', data[:4])
                win_org_x, win_org_y = x, y
            elif func == 0x020c and len(data) >= 4:  # SETWINDOWEXT
                h, w = struct.unpack('<hh', data[:4])
                if w != 0 and h != 0:
                    win_ext_w, win_ext_h = abs(w), abs(h)

        # Resolution scaling (300 DPI equivalent)
        if target_width_pt and target_height_pt:
            scale = 300.0 / 72.0
            img_w = max(20, int(target_width_pt * scale))
            img_h = max(20, int(target_height_pt * scale))
        else:
            scale_factor = (300.0 / float(inch)) if (inch and inch > 0) else 0.2
            img_w = max(20, int(win_ext_w * scale_factor))
            img_h = max(20, int(win_ext_h * scale_factor))

        img = PILImage.new('RGB', (img_w, img_h), (255, 255, 255))
        draw = ImageDraw.Draw(img)

        def tx(x):
            return int((x - win_org_x) * img_w / win_ext_w) if win_ext_w > 0 else 0

        def ty(y):
            return int((y - win_org_y) * img_h / win_ext_h) if win_ext_h > 0 else 0

        def decode_color(val):
            return (val & 0xff, (val >> 8) & 0xff, (val >> 16) & 0xff)

        objects = {}
        current_pen = {'color': (0, 0, 0), 'width': 1}
        current_brush = {'color': None}
        current_font = {'family': 'Times New Roman', 'size': 14, 'bold': False, 'italic': False}
        text_align = 0
        text_color = (0, 0, 0)
        cur_wmf_x, cur_wmf_y = 0, 0
        cur_pos = (0, 0)

        sym_map = {
            0x2d: '−', 0x2b: '+', 0x3d: '=', 0xb1: '±', 0xd7: '×', 0xf7: '÷', 0xb9: '≠',
            0xa3: '≤', 0xb3: '≥', 0xae: '→', 0xac: '←', 0xde: '⇒', 0xce: '∈', 0xcf: '∉',
            0xb0: '°', 0xa5: '∞', 0xb6: '∂', 0xd1: '∇', 0xf2: '∫',
            0x61: 'α', 0x62: 'β', 0x67: 'γ', 0x64: 'δ', 0x65: 'ε', 0x71: 'θ', 0x6c: 'λ',
            0x6d: 'μ', 0x70: 'π', 0x73: 'σ', 0x74: 'τ', 0x66: 'φ', 0x77: 'ω', 0x44: 'Δ',
            0xe6: '⎡', 0xe7: '⎢', 0xe8: '⎣', 0xf6: '⎤', 0xf7: '⎥', 0xf8: '⎦',
            0xe9: '⎛', 0xea: '⎜', 0xeb: '⎝', 0xf9: '⎞', 0xfa: '⎟', 0xfb: '⎠',
            0xec: '⎧', 0xed: '⎨', 0xee: '⎩', 0xef: '⎪', 0xfc: '⎫', 0xfd: '⎬', 0xfe: '⎭'
        }

        def _get_font(f_size):
            candidates = [
                'times.ttf',
                'DejaVuSerif.ttf',
                'LiberationSerif-Regular.ttf',
                '/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf',
                '/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf',
                '/usr/share/fonts/truetype/freefont/FreeSerif.ttf'
            ]
            for c in candidates:
                try:
                    return ImageFont.truetype(c, f_size)
                except Exception:
                    continue
            return ImageFont.load_default()

        for func, data in records:
            if func == 0x012e and len(data) >= 2:  # SETTEXTALIGN
                text_align = struct.unpack('<H', data[:2])[0]

            elif func == 0x02fb:  # CREATEFONTINDIRECT
                if len(data) >= 18:
                    h, w, esc, orient, weight, italic, under, strike, charset = struct.unpack('<hhhhhBBBB', data[:14])
                    facename = data[18:].split(b'\x00')[0].decode('latin1', errors='ignore')
                    f_size_px = int(abs(h) * img_h / win_ext_h) if win_ext_h > 0 and h != 0 else 14
                    font_obj = {
                        'type': 'font',
                        'family': facename or 'Times New Roman',
                        'size': max(10, f_size_px),
                        'bold': weight >= 600,
                        'italic': bool(italic)
                    }
                    slot = 0
                    while slot in objects: slot += 1
                    objects[slot] = font_obj

            elif func == 0x02fa:  # CREATEPENINDIRECT
                if len(data) >= 8:
                    style, w_x, w_y, color_val = struct.unpack('<HHHI', data[:10]) if len(data) >= 10 else (0, 1, 1, 0)
                    pen_w = max(1, int(w_x * img_w / win_ext_w)) if win_ext_w > 0 else 1
                    pen_obj = {
                        'type': 'pen',
                        'color': decode_color(color_val),
                        'width': pen_w
                    }
                    slot = 0
                    while slot in objects: slot += 1
                    objects[slot] = pen_obj

            elif func == 0x02fc:  # CREATEBRUSHINDIRECT
                if len(data) >= 8:
                    style, color_val, hatch = struct.unpack('<HIH', data[:8])
                    brush_obj = {
                        'type': 'brush',
                        'color': None if style == 1 else decode_color(color_val)
                    }
                    slot = 0
                    while slot in objects: slot += 1
                    objects[slot] = brush_obj

            elif func == 0x012d:  # SELECTOBJECT
                if len(data) >= 2:
                    idx = struct.unpack('<H', data[:2])[0]
                    obj = objects.get(idx)
                    if obj:
                        if obj['type'] == 'font': current_font = obj
                        elif obj['type'] == 'pen': current_pen = obj
                        elif obj['type'] == 'brush': current_brush = obj

            elif func == 0x01f0:  # DELETEOBJECT
                if len(data) >= 2:
                    idx = struct.unpack('<H', data[:2])[0]
                    objects.pop(idx, None)

            elif func == 0x0209:  # SETTEXTCOLOR
                if len(data) >= 4:
                    text_color = decode_color(struct.unpack('<I', data[:4])[0])

            elif func == 0x0214:  # MOVETO
                if len(data) >= 4:
                    y, x = struct.unpack('<hh', data[:4])
                    cur_wmf_x, cur_wmf_y = x, y
                    cur_pos = (tx(x), ty(y))

            elif func == 0x0213:  # LINETO
                if len(data) >= 4:
                    y, x = struct.unpack('<hh', data[:4])
                    target_p = (tx(x), ty(y))
                    draw.line([cur_pos, target_p], fill=current_pen['color'], width=current_pen['width'])
                    cur_pos = target_p
                    cur_wmf_x, cur_wmf_y = x, y

            elif func == 0x041b:  # RECTANGLE
                if len(data) >= 8:
                    b, r, t, l = struct.unpack('<hhhh', data[:8])
                    draw.rectangle([(tx(l), ty(t)), (tx(r), ty(b))], fill=current_brush['color'], outline=current_pen['color'], width=current_pen['width'])

            elif func == 0x0324:  # POLYGON
                if len(data) >= 2:
                    num_pts = struct.unpack('<h', data[:2])[0]
                    if len(data) >= 2 + num_pts * 4:
                        pts = [(tx(struct.unpack('<h', data[2 + i*4 : 4 + i*4])[0]), ty(struct.unpack('<h', data[4 + i*4 : 6 + i*4])[0])) for i in range(num_pts)]
                        draw.polygon(pts, fill=current_brush['color'], outline=current_pen['color'])

            elif func == 0x0325:  # POLYLINE
                if len(data) >= 2:
                    num_pts = struct.unpack('<h', data[:2])[0]
                    if len(data) >= 2 + num_pts * 4:
                        pts = [(tx(struct.unpack('<h', data[2 + i*4 : 4 + i*4])[0]), ty(struct.unpack('<h', data[4 + i*4 : 6 + i*4])[0])) for i in range(num_pts)]
                        draw.line(pts, fill=current_pen['color'], width=current_pen['width'])

            elif func == 0x0538:  # POLYPOLYGON
                if len(data) >= 2:
                    num_polys = struct.unpack('<h', data[:2])[0]
                    counts_offset = 2
                    pts_offset = counts_offset + num_polys * 2
                    if len(data) >= pts_offset:
                        poly_counts = struct.unpack(f'<{num_polys}h', data[counts_offset:pts_offset])
                        curr_pt_offset = pts_offset
                        for count in poly_counts:
                            if curr_pt_offset + count * 4 <= len(data):
                                pts = [(tx(struct.unpack('<h', data[curr_pt_offset + i*4 : curr_pt_offset + i*4 + 2])[0]),
                                        ty(struct.unpack('<h', data[curr_pt_offset + i*4 + 2 : curr_pt_offset + (i+1)*4])[0])) for i in range(count)]
                                curr_pt_offset += count * 4
                                draw.polygon(pts, fill=current_brush['color'], outline=current_pen['color'])

            elif func == 0x0a32:  # EXTTEXTOUT
                if len(data) >= 8:
                    y, x, count, options = struct.unpack('<hhhh', data[:8])
                    str_offset = 8
                    if (options & 6) != 0 and len(data) >= 16:
                        str_offset = 16
                    raw_bytes = data[str_offset : str_offset + count]

                    pad = count % 2
                    dx_offset = str_offset + count + pad
                    has_dx = len(data) >= dx_offset + count * 2

                    dx_array = []
                    if has_dx:
                        dx_array = struct.unpack(f'<{count}h', data[dx_offset : dx_offset + count * 2])

                    # Determine drawing origin (respect TA_UPDATECP or MOVETO position)
                    use_cp = bool(text_align & 1) or (x == 0 and y == 0 and (cur_wmf_x != 0 or cur_wmf_y != 0))
                    draw_wmf_x = cur_wmf_x if use_cp else x
                    draw_wmf_y = cur_wmf_y if use_cp else y

                    f_size = max(12, current_font['size'])
                    font = _get_font(f_size)
                    is_symbol = 'symbol' in current_font.get('family', '').lower()

                    cur_ch_wmf_x = draw_wmf_x
                    for i, b in enumerate(raw_bytes):
                        if is_symbol and b in sym_map:
                            ch_str = sym_map[b]
                        else:
                            ch_str = chr(b) if 32 <= b <= 126 else (chr(b) if b != 0 else '')
                        if ch_str:
                            draw.text((tx(cur_ch_wmf_x), ty(draw_wmf_y) - int(f_size * 0.8)), ch_str, font=font, fill=text_color)
                        if dx_array and i < len(dx_array):
                            cur_ch_wmf_x += dx_array[i]

                    cur_wmf_x = cur_ch_wmf_x
                    cur_pos = (tx(cur_wmf_x), ty(draw_wmf_y))

        return img
    except Exception as e:
        print(f"Warning: WMF parsing error: {e}")
        return None


def _process_equation_ole_image(part_blob, file_id, eq_filename, eq_dir, target_width_pt=None, target_height_pt=None):
    """
    Convert legacy Equation Editor 3.0 / MathType image (WMF/EMF/PNG) to a crisp, high-resolution PNG asset.
    Uses Windows GDI+ on Windows, and a pure-Python WMF rasterizer on Linux (Render).
    """
    os.makedirs(eq_dir, exist_ok=True)
    out_path = os.path.join(eq_dir, eq_filename)
    if os.path.exists(out_path):
        try:
            with PILImage.open(out_path) as img:
                return out_path, img.width, img.height
        except Exception:
            pass

    # Method 1: High-Fidelity Windows GDI+ Rasterization (if running on Windows)
    if _gdiplus_available:
        try:
            temp_wmf = out_path + '.tmp.wmf'
            with open(temp_wmf, 'wb') as f:
                f.write(part_blob)

            image = ctypes.c_void_p()
            st = _gdiplus.GdipLoadImageFromFile(ctypes.c_wchar_p(temp_wmf), ctypes.byref(image))
            if st == 0:
                width = ctypes.c_float()
                height = ctypes.c_float()
                _gdiplus.GdipGetImageDimension(image, ctypes.byref(width), ctypes.byref(height))

                scale = 300.0 / 72.0
                if target_width_pt and target_height_pt:
                    w_px = max(10, int(target_width_pt * scale))
                    h_px = max(10, int(target_height_pt * scale))
                else:
                    ar = (width.value / height.value) if height.value > 0 else 1.0
                    h_px = int(24.0 * scale)
                    w_px = max(10, int(h_px * ar))

                bitmap = ctypes.c_void_p()
                _gdiplus.GdipCreateBitmapFromScan0(w_px, h_px, 0, 0x22009, None, ctypes.byref(bitmap))

                graphics = ctypes.c_void_p()
                _gdiplus.GdipGetImageGraphicsContext(bitmap, ctypes.byref(graphics))

                # Pure white background, anti-aliased
                _gdiplus.GdipGraphicsClear(graphics, ctypes.c_uint(0xFFFFFFFF))
                _gdiplus.GdipSetSmoothingMode(graphics, 2)
                _gdiplus.GdipSetInterpolationMode(graphics, 7)
                _gdiplus.GdipSetTextRenderingHint(graphics, 4)

                _gdiplus.GdipDrawImageRectRect(
                    graphics, image,
                    0.0, 0.0, float(w_px), float(h_px),
                    0.0, 0.0, width.value, height.value,
                    2, None, None, None
                )

                _gdiplus.GdipSaveImageToFile(bitmap, ctypes.c_wchar_p(out_path), ctypes.byref(_png_clsid), None)
                _gdiplus.GdipDeleteGraphics(graphics)
                _gdiplus.GdipDisposeImage(bitmap)
                _gdiplus.GdipDisposeImage(image)

                if os.path.exists(temp_wmf):
                    os.remove(temp_wmf)

                # Auto-trim excess white borders
                with PILImage.open(out_path) as rendered_img:
                    trimmed = _trim_white_borders(rendered_img)
                    trimmed.save(out_path, format='PNG')
                    return out_path, trimmed.width, trimmed.height
        except Exception as e:
            print(f"Warning: GDI+ WMF render exception: {e}")
            if os.path.exists(temp_wmf):
                try: os.remove(temp_wmf)
                except Exception: pass

    # Method 2: Pure-Python WMF Rasterizer (Cross-platform, Works on Linux / Render)
    try:
        wmf_img = _render_wmf_to_pil(part_blob, target_width_pt=target_width_pt, target_height_pt=target_height_pt)
        if wmf_img:
            trimmed = _trim_white_borders(wmf_img)
            trimmed.save(out_path, format='PNG')
            return out_path, trimmed.width, trimmed.height
    except Exception as e:
        print(f"Warning: Pure Python WMF render exception: {e}")

    # Method 3: Standard Pillow Fallback (for embedded PNG, JPEG, GIF, BMP)
    try:
        img = PILImage.open(BytesIO(part_blob))
        if img.mode != 'RGB':
            img = img.convert('RGB')
        trimmed = _trim_white_borders(img)
        trimmed.save(out_path, format='PNG')
        return out_path, trimmed.width, trimmed.height
    except Exception as e:
        print(f"Warning: Failed to convert legacy OLE equation image: {e}")
        return None, 0, 0


def omml_to_latex(elem):
    """
    Convert Word OMML math XML element to standard LaTeX string.
    Supports matrices, delimiters, fractions, superscripts, subscripts, radicals,
    n-ary sums/integrals/products, accents, over/under braces, functions, and equation arrays.
    """
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    
    if tag in ('oMath', 'oMathPara', 'e', 'num', 'den', 'sub', 'sup', 'fName', 'lim', 'deg', 'limLow', 'limUpp'):
        return ''.join(omml_to_latex(child) for child in elem)
    
    elif tag == 'm':
        # Matrix representation
        rows = []
        for mr in elem.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}mr'):
            cells = [omml_to_latex(e).strip() for e in mr.findall('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')]
            rows.append(' & '.join(cells))
        matrix_body = ' \\\\ '.join(rows)
        return f'\\begin{{matrix}} {matrix_body} \\end{{matrix}}'

    elif tag == 'eqArr':
        # Equation array / aligned equations
        lines = []
        for e in elem.findall('{http://schemas.openxmlformats.org/officeDocument/2006/math}e'):
            lines.append(omml_to_latex(e).strip())
        body = ' \\\\ '.join(lines)
        return f'\\begin{{matrix}} {body} \\end{{matrix}}'

    elif tag == 'r':
        text = ''
        for child in elem:
            ctag = child.tag.split('}')[-1]
            if ctag == 't':
                text += child.text or ''
        replacements = {
            '…': r'\dots ', '⋯': r'\cdots ', '⋮': r'\vdots ', '⋱': r'\ddots ',
            '∞': r'\infty ', '∑': r'\sum ', 'π': r'\pi ', '∏': r'\prod ',
            'α': r'\alpha ', 'β': r'\beta ', 'γ': r'\gamma ', 'δ': r'\delta ', 'ε': r'\epsilon ', 'ϵ': r'\varepsilon ',
            'θ': r'\theta ', 'ϑ': r'\vartheta ', 'λ': r'\lambda ', 'μ': r'\mu ', 'σ': r'\sigma ', 'τ': r'\tau ',
            'φ': r'\phi ', 'ϕ': r'\varphi ', 'ω': r'\omega ', 'η': r'\eta ', 'κ': r'\kappa ', 'ρ': r'\rho ',
            'ψ': r'\psi ', 'ξ': r'\xi ', 'ζ': r'\zeta ',
            'Δ': r'\Delta ', 'Σ': r'\Sigma ', 'Ω': r'\Omega ', 'Γ': r'\Gamma ', 'Λ': r'\Lambda ',
            'Φ': r'\Phi ', 'Ψ': r'\Psi ', 'Θ': r'\Theta ', 'Π': r'\Pi ',
            '≤': r'\le ', '≥': r'\ge ', '≠': r'\ne ', '×': r'\times ', '÷': r'\div ', '±': r'\pm ', '∓': r'\mp ',
            '≈': r'\approx ', '≡': r'\equiv ', '∼': r'\sim ', '∝': r'\propto ',
            '∈': r'\in ', '∉': r'\notin ', '⊂': r'\subset ', '⊆': r'\subseteq ', '∪': r'\cup ', '∩': r'\cap ',
            '→': r'\to ', '←': r'\gets ', '⇒': r'\Rightarrow ', '⇔': r'\Leftrightarrow ', '↔': r'\leftrightarrow ',
            '∂': r'\partial ', '∇': r'\nabla ', '∫': r'\int ', '∬': r'\iint ', '∭': r'\iiint ', '∮': r'\oint ',
            '·': r'\cdot ', '°': r'^\circ ', '′': "'", '″': "''"
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text

    elif tag == 'f':
        num_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}num')
        den_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}den')
        num = omml_to_latex(num_elem).strip() if num_elem is not None else ''
        den = omml_to_latex(den_elem).strip() if den_elem is not None else ''
        return f'\\frac{{{num}}}{{{den}}}'

    elif tag == 'sSup':
        e_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')
        sup_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}sup')
        e_str = omml_to_latex(e_elem).strip() if e_elem is not None else ''
        sup_str = omml_to_latex(sup_elem).strip() if sup_elem is not None else ''
        return f'{{{e_str}}}^{{{sup_str}}}'

    elif tag == 'sSub':
        e_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')
        sub_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}sub')
        e_str = omml_to_latex(e_elem).strip() if e_elem is not None else ''
        sub_str = omml_to_latex(sub_elem).strip() if sub_elem is not None else ''
        return f'{{{e_str}}}_{{{sub_str}}}'

    elif tag == 'sSubSup':
        e_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')
        sub_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}sub')
        sup_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}sup')
        e_str = omml_to_latex(e_elem).strip() if e_elem is not None else ''
        sub_str = omml_to_latex(sub_elem).strip() if sub_elem is not None else ''
        sup_str = omml_to_latex(sup_elem).strip() if sup_elem is not None else ''
        return f'{{{e_str}}}_{{{sub_str}}}^{{{sup_str}}}'

    elif tag == 'd':
        dPr = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}dPr')
        beg_chr = '('
        end_chr = ')'
        if dPr is not None:
            beg = dPr.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}begChr')
            end = dPr.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}endChr')
            if beg is not None: beg_chr = beg.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', beg_chr)
            if end is not None: end_chr = end.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', end_chr)
        e_elems = elem.findall('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')
        inner = ''.join(omml_to_latex(e) for e in e_elems)
        beg_map = {'(': r'\left(', '[': r'\left[', '{': r'\left\{', '|': r'\left|', '‖': r'\left\|', '': r'\left.'}
        end_map = {')': r'\right)', ']': r'\right]', '}': r'\right\}', '|': r'\right|', '‖': r'\right\|', '': r'\right.'}
        b_str = beg_map.get(beg_chr, f'\\left{beg_chr}' if beg_chr else r'\left.')
        e_str = end_map.get(end_chr, f'\\right{end_chr}' if end_chr else r'\right.')
        return f'{b_str}{inner}{e_str}'

    elif tag == 'nary':
        naryPr = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}naryPr')
        chr_val = '∑'
        if naryPr is not None:
            c = naryPr.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}chr')
            if c is not None: chr_val = c.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', chr_val)
        sub_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}sub')
        sup_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}sup')
        e_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')
        
        sub_txt = omml_to_latex(sub_elem).strip() if sub_elem is not None else ''
        sup_txt = omml_to_latex(sup_elem).strip() if sup_elem is not None else ''
        sub_str = f'_{{{sub_txt}}}' if sub_txt else ''
        sup_str = f'^{{{sup_txt}}}' if sup_txt else ''
        e_str = omml_to_latex(e_elem).strip() if e_elem is not None else ''
        
        op_symbol = r'\sum' if chr_val in ('∑', '\u2211') else (
            r'\int' if chr_val in ('∫', '\u222b') else (
                r'\iint' if chr_val in ('∬', '\u222c') else (
                    r'\iiint' if chr_val in ('∭', '\u222d') else (
                        r'\oint' if chr_val in ('∮', '\u222e') else (
                            r'\prod' if chr_val in ('∏', '\u220f') else chr_val
                        )
                    )
                )
            )
        )
        return f'{op_symbol}{sub_str}{sup_str}{{{e_str}}}'

    elif tag == 'func':
        fname_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}fName')
        e_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')
        fname_str = omml_to_latex(fname_elem).strip() if fname_elem is not None else ''
        e_str = omml_to_latex(e_elem).strip() if e_elem is not None else ''
        std_funcs = {'sin', 'cos', 'tan', 'sec', 'csc', 'cot', 'sinh', 'cosh', 'tanh', 'ln', 'log', 'exp', 'lim', 'min', 'max', 'det', 'gcd', 'deg', 'dim', 'hom', 'ker', 'arg'}
        fn_clean = fname_str.strip().replace('\\', '')
        if fn_clean in std_funcs:
            return f'\\{fn_clean}{{{e_str}}}'
        else:
            return f'{fname_str}{{{e_str}}}'

    elif tag == 'rad':
        deg_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}deg')
        e_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')
        deg_txt = omml_to_latex(deg_elem).strip() if deg_elem is not None else ''
        deg_str = f'[{deg_txt}]' if deg_txt else ''
        e_str = omml_to_latex(e_elem).strip() if e_elem is not None else ''
        return f'\\sqrt{deg_str}{{{e_str}}}'

    elif tag == 'groupChr':
        groupChrPr = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}groupChrPr')
        pos_val = 'top'
        if groupChrPr is not None:
            p = groupChrPr.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}pos')
            if p is not None: pos_val = p.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', pos_val)
        e_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')
        e_str = omml_to_latex(e_elem).strip() if e_elem is not None else ''
        if pos_val == 'top':
            return f'\\overbrace{{{e_str}}}'
        else:
            return f'\\underbrace{{{e_str}}}'

    elif tag == 'bar':
        e_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')
        e_str = omml_to_latex(e_elem).strip() if e_elem is not None else ''
        return f'\\overline{{{e_str}}}'

    elif tag == 'acc':
        accPr = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}accPr')
        chr_val = '^'
        if accPr is not None:
            c = accPr.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}chr')
            if c is not None: chr_val = c.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', chr_val)
        e_elem = elem.find('{http://schemas.openxmlformats.org/officeDocument/2006/math}e')
        e_str = omml_to_latex(e_elem).strip() if e_elem is not None else ''
        acc_map = {'^': r'\hat', '¯': r'\bar', '⃗': r'\vec', '˙': r'\dot', '¨': r'\ddot', '~': r'\tilde'}
        acc_cmd = acc_map.get(chr_val, r'\hat')
        return f'{acc_cmd}{{{e_str}}}'

    else:
        return ''.join(omml_to_latex(child) for child in elem)


def _render_equation_asset(latex_str, file_id, eq_filename, eq_dir):
    """Render LaTeX string to a crisp, high-resolution PNG asset with pure white background."""
    os.makedirs(eq_dir, exist_ok=True)
    out_path = os.path.join(eq_dir, eq_filename)
    if os.path.exists(out_path):
        try:
            with PILImage.open(out_path) as img:
                return out_path, img.width, img.height
        except Exception:
            pass

    try:
        fig = plt.figure(figsize=(0.1, 0.1), facecolor='white')
        formatted_latex = f"${latex_str}$" if not latex_str.startswith('$') else latex_str
        text = fig.text(0, 0, formatted_latex, fontsize=11, color='black')
        fig.canvas.draw()
        bbox = text.get_window_extent(fig.canvas.get_renderer())
        w_in, h_in = bbox.width / fig.dpi, bbox.height / fig.dpi
        fig.set_size_inches(w_in + 0.04, h_in + 0.04)
        plt.savefig(out_path, dpi=300, bbox_inches='tight', facecolor='white', pad_inches=0.01)
        plt.close(fig)

        with PILImage.open(out_path) as img:
            trimmed = _trim_white_borders(img)
            trimmed.save(out_path, format='PNG')
            return out_path, trimmed.width, trimmed.height
    except Exception as e:
        print(f"Warning: Failed to render equation latex '{latex_str}': {e}")
        return None, 0, 0


def _parse_cell_content(cell, doc, file_id, eq_dir, row_idx):
    """
    Parse a cell's paragraphs to extract structured content segments (text runs, OMML equations, and legacy OLE equations).
    Maintains exact inline position and document order.
    Returns tuple: (full_text_string, content_list)
    """
    content = []
    current_text = ""
    eq_counter = 0

    for p in cell.paragraphs:
        for child in p._element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag in ('oMath', 'oMathPara'):
                if current_text:
                    content.append({'type': 'text', 'value': current_text})
                    current_text = ""

                try:
                    omml_raw = ET.tostring(child, encoding='utf-8').decode('utf-8')
                    omml_xml_clean = re.sub(r'xmlns:ns\d+="[^"]*"', '', omml_raw)
                    omml_xml_clean = re.sub(r'ns\d+:', 'm:', omml_xml_clean)
                    if 'xmlns:m=' not in omml_xml_clean:
                        omml_xml_clean = re.sub(
                            r'^<m:(oMathPara|oMath)\b',
                            r'<m:\1 xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
                            omml_xml_clean
                        )

                    latex = omml_to_latex(child)
                    if latex.strip():
                        eq_counter += 1
                        eq_filename = f"eq_{row_idx}_{eq_counter}.png"
                        out_path, w, h = _render_equation_asset(latex, file_id, eq_filename, eq_dir)
                        url = f"/api/equations/{file_id}/{eq_filename}"

                        # Matplotlib renders at 300 DPI: 1 inch = 72 pt = 300 px
                        w_pt = float(w) * (72.0 / 300.0) if w > 0 else 20.0
                        h_pt = float(h) * (72.0 / 300.0) if h > 0 else 14.0
                        aspect_ratio = (w_pt / h_pt) if h_pt > 0 else 1.0
                        is_block = (tag == 'oMathPara') or (h_pt >= 22.0) or ('\\begin{matrix}' in latex)

                        content.append({
                            'type': 'equation',
                            'latex': latex,
                            'omml': omml_xml_clean,
                            'local_path': out_path,
                            'url': url,
                            'width': w,
                            'height': h,
                            'width_pt': w_pt,
                            'height_pt': h_pt,
                            'aspect_ratio': aspect_ratio,
                            'is_legacy_ole': False,
                            'is_block': is_block
                        })
                except Exception as e:
                    print(f"Warning: Exception parsing OMML equation in row {row_idx}: {e}")

            elif tag in ('object', 'pict'):
                # Direct OLE Object or legacy picture
                if current_text:
                    content.append({'type': 'text', 'value': current_text})
                    current_text = ""

                eq_obj = _extract_ole_equation_item(child, doc, file_id, eq_dir, row_idx, eq_counter + 1)
                if eq_obj:
                    eq_counter += 1
                    content.append(eq_obj)

            elif tag == 'r':
                # Walk run child elements
                for r_child in child:
                    r_tag = r_child.tag.split('}')[-1] if '}' in r_child.tag else r_child.tag
                    if r_tag == 't':
                        current_text += (r_child.text or '')
                    elif r_tag in ('object', 'pict'):
                        if current_text:
                            content.append({'type': 'text', 'value': current_text})
                            current_text = ""
                        eq_obj = _extract_ole_equation_item(r_child, doc, file_id, eq_dir, row_idx, eq_counter + 1)
                        if eq_obj:
                            eq_counter += 1
                            content.append(eq_obj)
                    elif r_tag in ('oMath', 'oMathPara'):
                        if current_text:
                            content.append({'type': 'text', 'value': current_text})
                            current_text = ""
                        try:
                            omml_raw = ET.tostring(r_child, encoding='utf-8').decode('utf-8')
                            omml_xml_clean = re.sub(r'xmlns:ns\d+="[^"]*"', '', omml_raw)
                            omml_xml_clean = re.sub(r'ns\d+:', 'm:', omml_xml_clean)
                            if 'xmlns:m=' not in omml_xml_clean:
                                omml_xml_clean = re.sub(
                                    r'^<m:(oMathPara|oMath)\b',
                                    r'<m:\1 xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
                                    omml_xml_clean
                                )

                            latex = omml_to_latex(r_child)
                            if latex.strip():
                                eq_counter += 1
                                eq_filename = f"eq_{row_idx}_{eq_counter}.png"
                                out_path, w, h = _render_equation_asset(latex, file_id, eq_filename, eq_dir)
                                url = f"/api/equations/{file_id}/{eq_filename}"
                                w_pt = float(w) * (72.0 / 300.0) if w > 0 else 20.0
                                h_pt = float(h) * (72.0 / 300.0) if h > 0 else 14.0
                                aspect_ratio = (w_pt / h_pt) if h_pt > 0 else 1.0
                                is_block = (r_tag == 'oMathPara') or (h_pt >= 22.0) or ('\\begin{matrix}' in latex)
                                content.append({
                                    'type': 'equation',
                                    'latex': latex,
                                    'omml': omml_xml_clean,
                                    'local_path': out_path,
                                    'url': url,
                                    'width': w,
                                    'height': h,
                                    'width_pt': w_pt,
                                    'height_pt': h_pt,
                                    'aspect_ratio': aspect_ratio,
                                    'is_legacy_ole': False,
                                    'is_block': is_block
                                })
                        except Exception as e:
                            print(f"Warning: Exception parsing run OMML equation in row {row_idx}: {e}")

        if current_text and not current_text.endswith(' '):
            current_text += ' '

    if current_text:
        content.append({'type': 'text', 'value': current_text})

    # Consolidate adjacent text segments in content
    consolidated = []
    for item in content:
        if consolidated and consolidated[-1]['type'] == 'text' and item['type'] == 'text':
            consolidated[-1]['value'] += item['value']
        else:
            consolidated.append(item)

    full_text = "".join([
        item['value'] if item['type'] == 'text' else (
            f" ${item.get('latex', '')}$ " if item.get('latex') else " [Equation] "
        ) for item in consolidated
    ]).strip()

    if not full_text:
        full_text = cell.text.strip()
        if not full_text and any(c['type'] == 'equation' for c in consolidated):
            full_text = "[Equation]"

    return full_text, consolidated


def _extract_ole_equation_item(obj_elem, doc, file_id, eq_dir, row_idx, eq_number):
    """Extract legacy Equation Editor 3.0 / MathType object as an equation content item."""
    imgd = obj_elem.find('.//{urn:schemas-microsoft-com:vml}imagedata')
    shape = obj_elem.find('.//{urn:schemas-microsoft-com:vml}shape')
    if imgd is None:
        return None

    rid = imgd.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
    if not rid or rid not in doc.part.rels:
        return None

    part = doc.part.rels[rid].target_part
    style = shape.attrib.get('style', '') if shape is not None else ''
    w_pt, h_pt = _parse_shape_style_dimensions(style)

    eq_filename = f"eq_ole_{row_idx}_{eq_number}.png"
    out_path, img_w, img_h = _process_equation_ole_image(part.blob, file_id, eq_filename, eq_dir, target_width_pt=w_pt, target_height_pt=h_pt)

    if not out_path:
        return None

    w_pt = w_pt or (float(img_w) * (72.0 / 300.0) if img_w > 0 else 40.0)
    h_pt = h_pt or (float(img_h) * (72.0 / 300.0) if img_h > 0 else 18.0)
    aspect_ratio = (w_pt / h_pt) if (h_pt > 0) else ((img_w / float(img_h)) if img_h > 0 else 1.0)
    url = f"/api/equations/{file_id}/{eq_filename}"
    is_block = (h_pt >= 22.0)

    return {
        'type': 'equation',
        'is_legacy_ole': True,
        'url': url,
        'local_path': out_path,
        'width': img_w,
        'height': img_h,
        'width_pt': w_pt,
        'height_pt': h_pt,
        'aspect_ratio': aspect_ratio,
        'is_block': is_block,
        'latex': ''
    }


def _parse_part_a(table, start_idx, end_idx, doc, file_id, img_dir, eq_dir=None):
    """Parse Part A questions into dictionary grouped by question number."""
    questions = {}
    if start_idx is None:
        return questions

    last_q = None

    for r_idx in range(start_idx + 1, end_idx):
        row = table.rows[r_idx]
        cells_text = [cell.text.strip() for cell in row.cells]
        if len(cells_text) < 3:
            continue

        eq_rids = _extract_row_equation_rids(row)
        row_imgs = _extract_row_images(row, r_idx, doc, file_id, img_dir, eq_rids)

        q_no_raw = cells_text[0].strip()
        alt_raw = cells_text[1].strip() if len(cells_text) > 1 else '1'
        
        # Extract rich text and equations from cell 2
        text, content = _parse_cell_content(row.cells[2], doc, file_id, eq_dir, r_idx) if eq_dir else (cells_text[2].strip(), [{'type': 'text', 'value': cells_text[2].strip()}])
        if not text and len(cells_text) > 2:
            text = cells_text[2].strip()

        k_level = cells_text[3].strip() if len(cells_text) > 3 else ''
        co = cells_text[4].strip() if len(cells_text) > 4 else ''

        q_match = re.search(r'(\d+)', q_no_raw)
        if not q_match:
            if row_imgs and last_q:
                for img_url in row_imgs:
                    if img_url not in last_q['images']:
                        last_q['images'].append(img_url)
            continue

        has_equations = any(c.get('type') == 'equation' for c in content)
        if not text and not row_imgs and not has_equations:
            continue

        q_no = int(q_match.group(1))
        alt_idx = int(alt_raw) if alt_raw.isdigit() else 1

        if q_no not in questions:
            questions[q_no] = []

        q_obj = {
            'text': text,
            'content': content,
            'k_level': k_level,
            'co': co,
            'alt_index': alt_idx,
            'images': list(row_imgs),
        }
        questions[q_no].append(q_obj)
        last_q = q_obj

    return questions


def _parse_part_bc(table, start_idx, end_idx, doc, file_id, img_dir, eq_dir=None):
    """Parse Part B or Part C questions into dictionary grouped by question number and sub-part (a/b)."""
    questions = {}
    if start_idx is None:
        return questions

    current_q_no = None
    last_q_sub = None

    for r_idx in range(start_idx + 1, end_idx):
        row = table.rows[r_idx]
        cells_text = [cell.text.strip() for cell in row.cells]
        if len(cells_text) < 3:
            continue

        eq_rids = _extract_row_equation_rids(row)
        row_imgs = _extract_row_images(row, r_idx, doc, file_id, img_dir, eq_rids)

        q_no_raw = cells_text[0].strip()
        alt_raw = cells_text[1].strip() if len(cells_text) > 1 else '1'
        
        # Extract rich text and equations from cell 2
        text, content = _parse_cell_content(row.cells[2], doc, file_id, eq_dir, r_idx) if eq_dir else (cells_text[2].strip(), [{'type': 'text', 'value': cells_text[2].strip()}])
        if not text and len(cells_text) > 2:
            text = cells_text[2].strip()

        k_level = cells_text[3].strip() if len(cells_text) > 3 else ''
        co = cells_text[4].strip() if len(cells_text) > 4 else ''
        sub_part_raw = cells_text[5].strip() if len(cells_text) > 5 else ''

        q_match = re.search(r'(\d+)', q_no_raw)
        if q_match:
            current_q_no = int(q_match.group(1))

        if not current_q_no:
            if row_imgs and last_q_sub:
                for img_url in row_imgs:
                    if img_url not in last_q_sub['images']:
                        last_q_sub['images'].append(img_url)
            continue

        has_equations = any(c.get('type') == 'equation' for c in content)
        if not text and not row_imgs and not has_equations:
            continue

        sub_match = re.search(r'([abAB])', q_no_raw + ' ' + sub_part_raw)
        sub_part = sub_match.group(1).lower() if sub_match else 'a'
        alt_idx = int(alt_raw) if alt_raw.isdigit() else 1

        if current_q_no not in questions:
            questions[current_q_no] = {'a': [], 'b': []}

        q_obj = {
            'text': text,
            'content': content,
            'k_level': k_level,
            'co': co,
            'alt_index': alt_idx,
            'images': list(row_imgs),
        }
        questions[current_q_no][sub_part].append(q_obj)
        last_q_sub = q_obj

    return questions
