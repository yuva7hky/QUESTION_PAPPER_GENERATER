"""
docx_generator.py — Generate professional CIE-I / CIE-II / Model Question Paper DOCX.

Uses python-docx to create a Word document matching the official
JIT examination paper format (borderless clean text layout matching PDF).
"""

import os
import uuid
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image as PILImage

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Column widths matching PDF A4 layout (Total = 18.0 cm)
QUESTION_COL_WIDTHS = [Cm(1.5), Cm(12.2), Cm(1.4), Cm(1.4), Cm(1.5)]


def _resolve_local_path(img_url, base_dir=BASE_DIR):
    """Convert relative or absolute URL /api/images/<file_id>/<filename> to local disk path."""
    if not img_url:
        return None
    if '/api/images/' in img_url:
        rel_path = img_url.split('/api/images/')[-1].lstrip('/')
        return os.path.join(base_dir, 'uploads', 'images', rel_path)
    return img_url


def _insert_question_images(cell, images, base_dir=BASE_DIR, max_cm=11.0):
    """Insert images associated with a question into a Word table cell."""
    if not images:
        return
    for img_url in images:
        local_path = _resolve_local_path(img_url, base_dir)
        if local_path and os.path.exists(local_path):
            try:
                with PILImage.open(local_path) as pil_img:
                    w, h = pil_img.size
                if w > 0:
                    orig_w_cm = w / 37.795
                    target_w_cm = min(max_cm, orig_w_cm)
                    target_w_cm = max(2.5, target_w_cm)
                else:
                    target_w_cm = 8.0
                
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run()
                run.add_picture(local_path, width=Cm(target_w_cm))
            except Exception as e:
                print(f"Warning: Could not insert image {local_path} into DOCX: {e}")


def generate_docx(paper_data, output_path=None):
    """
    Generate a DOCX question paper matching the printed JIT exam paper and PDF export.
    
    Args:
        paper_data: Complete paper dict from generator.generate_paper()
        output_path: Path to write the DOCX file
    """
    if not output_path:
        out_dir = os.path.join(BASE_DIR, 'generated')
        os.makedirs(out_dir, exist_ok=True)
        paper_id = paper_data.get('paper_id', str(uuid.uuid4())[:8])
        filename = f"Question_Paper_{paper_id}.docx"
        output_path = os.path.join(out_dir, filename)

    doc = Document()

    # ── Page Setup (Margins match PDF) ────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    metadata = paper_data.get('metadata', {})
    course_outcomes = paper_data.get('course_outcomes', [])
    part_a = paper_data.get('part_a', {})
    part_b = paper_data.get('part_b', {})
    part_c = paper_data.get('part_c', {})

    # ── Register Number Box (Top Right) ───────────────────
    reg_table = doc.add_table(rows=1, cols=13)
    reg_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    reg_table.autofit = False
    _remove_table_borders(reg_table)
    
    reg_cell_0 = reg_table.cell(0, 0)
    reg_cell_0.width = Cm(1.8)
    reg_cell_0.paragraphs[0].text = "Reg No"
    _set_font(reg_cell_0.paragraphs[0].runs[0], bold=True, size=11)
    # Prevent "Reg No" from wrapping to two lines
    tcPr = reg_cell_0._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(r'<w:noWrap %s/>' % nsdecls('w')))

    for i in range(1, 13):
        cell = reg_table.cell(0, i)
        _set_single_cell_border(cell)
        cell.width = Cm(0.5)

    _add_spacer(doc, 2)

    # ── College Header ────────────────────────────────────
    _add_paragraph(doc, "JEPPIAAR INSTITUTE OF TECHNOLOGY", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "(An Autonomous Institution)", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, '"Self-Belief | Self-Discipline | Self-Respect"', italic=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "Kunnam, Sunguvarchatram, Sriperumbudur – 631 604.", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_spacer(doc, 4)

    # ── Exam Title ────────────────────────────────────────
    exam_type = metadata.get('exam_type', 'CIE I').upper()
    month_year = metadata.get('month_year', '').upper()
    _add_paragraph(doc, f"{exam_type} – {month_year}", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_spacer(doc, 4)

    # ── Metadata Grid ─────────────────────────────────────
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(meta_table)

    subject_code = metadata.get('su00') if metadata.get('su00') and metadata.get('su00') != '-' else metadata.get('subject_code', '-')
    subject_name = metadata.get('su01') if metadata.get('su01') and metadata.get('su01') != '-' else metadata.get('subject_name', '-')
    branch_info = metadata.get('branch_info') if metadata.get('branch_info') and metadata.get('branch_info') != '-' else '-'
    duration = metadata.get('duration', '1 ½ hours')
    date = metadata.get('date', '___________')
    max_marks = metadata.get('max_marks_display', metadata.get('max_marks', '-'))

    _set_meta_row(meta_table, 0, f"SUB CODE: {subject_code}", f"SUBJECT: {subject_name}")
    _set_meta_row(meta_table, 1, f"Duration: {duration}", f"Branch / Year / Sem: {branch_info}")
    _set_meta_row(meta_table, 2, f"Date: {date}", f"Maximum: {max_marks}")

    _set_meta_table_widths(meta_table)
    _add_spacer(doc, 4)

    # ── Course Outcomes ───────────────────────────────────
    if course_outcomes:
        p = doc.add_paragraph()
        run = p.add_run("Course Outcome: – After Successful Completion of the Course, the Students should be able to")
        run.bold = True
        _set_font(run, size=11)

        for co in course_outcomes:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            run1 = p.add_run(f"{co['id']}    ")
            run1.bold = True
            _set_font(run1, size=11)
            run2 = p.add_run(co['text'])
            _set_font(run2, size=11)

    _add_spacer(doc, 4)

    # ── Instructions ──────────────────────────────────────
    _add_paragraph(doc, "Answer all questions.", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_spacer(doc, 4)

    # ═══════════ PART A ═══════════════════════════════════
    if part_a and part_a.get('questions'):
        config = part_a.get('config', '')
        _add_paragraph(doc, f"PART – A ({config})", bold=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_spacer(doc, 2)

        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)
        
        _set_table_headers(table)

        for q in part_a['questions']:
            row_cells = table.add_row().cells
            _set_cell_text(row_cells[0], str(q['q_no']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_question_content(row_cells[1], q)
            _insert_question_images(row_cells[1], q.get('images', []))
            _set_cell_text(row_cells[2], str(q['co']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row_cells[3], str(q['marks']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row_cells[4], str(q['k_level']), align=WD_ALIGN_PARAGRAPH.CENTER)

        _set_table_col_widths(table, QUESTION_COL_WIDTHS)
        _add_spacer(doc, 6)

    # ═══════════ PART B ═══════════════════════════════════
    if part_b and part_b.get('questions'):
        config = part_b.get('config', '')
        _add_paragraph(doc, f"PART – B ({config})", bold=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_spacer(doc, 2)

        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        _set_table_headers(table)

        for group in part_b['questions']:
            q_no = group['q_no']
            a_q = group['a']
            b_q = group['b']

            # (a)
            row1 = table.add_row().cells
            _set_cell_text(row1[0], f"{q_no}.", align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_question_content(row1[1], a_q, prefix="a) ")
            _insert_question_images(row1[1], a_q.get('images', []))
            _set_cell_text(row1[2], str(a_q['co']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row1[3], str(a_q['marks']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row1[4], str(a_q['k_level']), align=WD_ALIGN_PARAGRAPH.CENTER)

            # (OR)
            row_or = table.add_row().cells
            _set_cell_text(row_or[1], "(OR)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

            # (b)
            row2 = table.add_row().cells
            _set_cell_question_content(row2[1], b_q, prefix="b) ")
            _insert_question_images(row2[1], b_q.get('images', []))
            _set_cell_text(row2[2], str(b_q['co']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row2[3], str(b_q['marks']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row2[4], str(b_q['k_level']), align=WD_ALIGN_PARAGRAPH.CENTER)

        _set_table_col_widths(table, QUESTION_COL_WIDTHS)
        _add_spacer(doc, 6)

    # ═══════════ PART C ═══════════════════════════════════
    if part_c and part_c.get('questions'):
        config = part_c.get('config', '')
        _add_paragraph(doc, f"PART – C ({config})", bold=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_spacer(doc, 2)

        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        _set_table_headers(table)

        for group in part_c['questions']:
            q_no = group['q_no']
            a_q = group['a']
            b_q = group['b']

            row1 = table.add_row().cells
            _set_cell_text(row1[0], f"{q_no}.", align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_question_content(row1[1], a_q, prefix="a) ")
            _insert_question_images(row1[1], a_q.get('images', []))
            _set_cell_text(row1[2], str(a_q['co']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row1[3], str(a_q['marks']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row1[4], str(a_q['k_level']), align=WD_ALIGN_PARAGRAPH.CENTER)

            row_or = table.add_row().cells
            _set_cell_text(row_or[1], "(OR)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

            row2 = table.add_row().cells
            _set_cell_question_content(row2[1], b_q, prefix="b) ")
            _insert_question_images(row2[1], b_q.get('images', []))
            _set_cell_text(row2[2], str(b_q['co']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row2[3], str(b_q['marks']), align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row2[4], str(b_q['k_level']), align=WD_ALIGN_PARAGRAPH.CENTER)

        _set_table_col_widths(table, QUESTION_COL_WIDTHS)
        _add_spacer(doc, 6)

    # ── EXAMCELL Footer Block ─────────────────────────────
    _add_paragraph(doc, "EXAMCELL", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "Jeppiaar Institute of Technology (Autonomous)", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "Kunnam, Sunguvarchatram, Sriperumbudur – 631 604.", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_spacer(doc, 4)

    # ── K-Level Legend ────────────────────────────────────
    _add_paragraph(
        doc,
        "K1-Remembering, K2-Understanding, K3-Applying, K4-Analysing, K5-Evaluating, K6-Creating",
        size=7.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.save(output_path)
    return output_path


# ── Helper Functions ──────────────────────────────────────────

def _set_font(run, bold=False, italic=False, size=12):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _add_paragraph(doc, text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _set_font(run, bold, italic, size)
    return p


def _add_spacer(doc, points=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(points)
    p.paragraph_format.space_after = Pt(0)


def _set_meta_row(table, row_idx, left_text, right_text):
    cell_l = table.cell(row_idx, 0)
    cell_r = table.cell(row_idx, 1)

    p_l = cell_l.paragraphs[0]
    p_l.paragraph_format.space_after = Pt(1)
    run_l = p_l.add_run(left_text)
    _set_font(run_l, size=11)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.paragraph_format.space_after = Pt(1)
    run_r = p_r.add_run(right_text)
    _set_font(run_r, size=11)


def _set_meta_table_widths(table):
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(10.0)


def _set_table_headers(table):
    hdr_cells = table.rows[0].cells
    headers = ['Q.NO', 'QUESTIONS', 'CO NO.', 'MARKS', 'K LEVEL']
    for i, h in enumerate(headers):
        _set_cell_text(hdr_cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def _resolve_equation_path(eq_url, base_dir=BASE_DIR):
    """Convert relative or absolute URL /api/equations/<file_id>/<filename> to local disk path."""
    if not eq_url:
        return None
    if '/api/equations/' in eq_url:
        rel_path = eq_url.split('/api/equations/')[-1].lstrip('/')
        return os.path.join(base_dir, 'uploads', 'equations', rel_path)
    return eq_url


def _set_cell_question_content(cell, q_or_text, prefix="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """
    Fill a Word table cell with question content.
    - Inline equations (OMML or small image): embedded within the current text paragraph
    - Block equations (matrices, systems): placed in their OWN paragraph within the cell,
      with a short spacer before and after. Text continues in a new paragraph after the block.
    This prevents rows from collapsing and text from overlapping in the output.
    """
    def _new_para(c):
        """Add a new paragraph to cell with standard formatting."""
        pp = c.add_paragraph()
        pp.alignment = align
        pp.paragraph_format.space_after = Pt(1)
        pp.paragraph_format.space_before = Pt(1)
        return pp

    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)

    if prefix:
        run_p = p.add_run(prefix)
        _set_font(run_p, bold=True)

    if isinstance(q_or_text, dict) and q_or_text.get('content'):
        for item in q_or_text['content']:
            if item['type'] == 'text':
                run = p.add_run(item['value'])
                _set_font(run, bold=bold)

            elif item['type'] == 'equation':
                omml_xml = item.get('omml')
                is_block = item.get('is_block', False)
                inserted = False

                # ── OMML equations (native Word math, always inline-capable) ──
                if omml_xml and not item.get('is_legacy_ole'):
                    try:
                        omath_elem = parse_xml(omml_xml)
                        if is_block:
                            # Flush current paragraph; add block math on its own paragraph
                            p_block = _new_para(cell)
                            p_block.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            p_block._p.append(omath_elem)
                            # Start a fresh paragraph for any following text
                            p = _new_para(cell)
                        else:
                            p._p.append(omath_elem)
                        inserted = True
                    except Exception as e:
                        print(f"Warning: Could not parse OMML XML into DOCX: {e}")

                # ── Image equations (OLE WMF rasterized or OMML fallback) ──
                if not inserted:
                    local_path = item.get('local_path') or _resolve_equation_path(item.get('url'))
                    if local_path and os.path.exists(local_path):
                        try:
                            disp_w_pt = float(item.get('orig_w_pt') or item.get('width_pt') or 40.0)
                            disp_h_pt = float(item.get('orig_h_pt') or item.get('height_pt') or 14.0)
                            # Guard: limit to question column width (~260pt)
                            if disp_w_pt > 260.0:
                                scale = 260.0 / disp_w_pt
                                disp_w_pt = 260.0
                                disp_h_pt = disp_h_pt * scale

                            if is_block:
                                # Block equation: own paragraph, left-aligned
                                p_block = _new_para(cell)
                                p_block.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                p_block.paragraph_format.space_before = Pt(3)
                                p_block.paragraph_format.space_after = Pt(3)
                                run = p_block.add_run()
                                run.add_picture(local_path, width=Pt(disp_w_pt), height=Pt(disp_h_pt))
                                # New paragraph for any text that follows the block
                                p = _new_para(cell)
                            else:
                                # Inline equation: embed in current paragraph run
                                run = p.add_run()
                                run.add_picture(local_path, width=Pt(disp_w_pt), height=Pt(disp_h_pt))
                        except Exception as e:
                            print(f"Warning: Could not insert equation image into DOCX: {e}")
                            latex = item.get('latex', '')
                            run = p.add_run(f" ${latex}$ " if latex else " [Equation] ")
                            _set_font(run, bold=bold)
    else:
        text = q_or_text.get('text', '') if isinstance(q_or_text, dict) else str(q_or_text)
        run = p.add_run(text)
        _set_font(run, bold=bold)



def _set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    _set_font(run, bold=bold)


def _set_table_col_widths(table, widths):
    """Ensure explicit column widths across all table cells in Word."""
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = width


def _remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
    tblPr.append(tblBorders)


def _set_single_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>' % nsdecls('w'))
    tcPr.append(tcBorders)
